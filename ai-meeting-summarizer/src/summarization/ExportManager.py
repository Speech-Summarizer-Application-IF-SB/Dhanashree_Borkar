"""
Export Manager Module
Exports meeting summaries to various formats (PDF, Markdown, DOCX)
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional

# Prefer markdown2 if available, otherwise fall back to python-markdown or a minimal safe fallback.
try:
    import markdown2  # type: ignore
    MARKDOWN2_AVAILABLE = True
    MARKDOWN_AVAILABLE = False
except Exception:
    MARKDOWN2_AVAILABLE = False
    try:
        import markdown as _markdown  # type: ignore
        MARKDOWN_AVAILABLE = True
    except Exception:
        MARKDOWN_AVAILABLE = False

def convert_markdown_to_html(md_text: str) -> str:
    """
    Convert markdown text to HTML using available libraries or a safe fallback.
    """
    if MARKDOWN2_AVAILABLE:
        return markdown2.markdown(md_text, extras=['tables', 'fenced-code-blocks'])
    if MARKDOWN_AVAILABLE:
        # Use python-markdown if installed
        return _markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    # Minimal safe fallback: escape and wrap in <pre> so content is still viewable
    import html as _html
    return f"<pre>{_html.escape(md_text)}</pre>"

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Install with: pip install python-docx")


class ExportManager:
    """
    Handles exporting meeting summaries to different formats
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        self.output_dir = output_dir or Path("outputs/exports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_pdf(
        self,
        summary,  # MeetingSummary object
        meeting_name: str,
        page_size=letter
    ) -> Path:
        """
        Export summary as PDF with professional formatting
        
        Args:
            summary: MeetingSummary object
            meeting_name: Name of the meeting
            page_size: Page size (letter or A4)
        
        Returns:
            Path to generated PDF file
        """
        print("📄 Generating PDF...")
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"{meeting_name.replace(' ', '_')}_{timestamp}.pdf"
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(filename),
            pagesize=page_size,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#764ba2'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = styles['BodyText']
        normal_style.fontSize = 11
        normal_style.leading = 14
        
        # Title
        elements.append(Paragraph(f"🎙️ {meeting_name}", title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Metadata table
        metadata_data = [
            ['Date:', summary.timestamp.strftime('%A, %B %d, %Y at %I:%M %p')],
            ['Duration:', f'{summary.duration:.1f} minutes'],
        ]
        
        if summary.participants:
            metadata_data.append(['Participants:', ', '.join(summary.participants)])
        
        metadata_table = Table(metadata_data, colWidths=[1.5*inch, 5*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        
        elements.append(metadata_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Summary section
        elements.append(Paragraph("📋 Summary", heading_style))
        elements.append(Paragraph(summary.summary, normal_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Key Points
        if summary.key_points:
            elements.append(Paragraph("🔑 Key Points", heading_style))
            for point in summary.key_points:
                elements.append(Paragraph(f"• {point}", normal_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Action Items
        if summary.action_items:
            elements.append(Paragraph("✅ Action Items", heading_style))
            for item in summary.action_items:
                elements.append(Paragraph(f"☐ {item}", normal_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Decisions
        if summary.decisions:
            elements.append(Paragraph("⚖️ Decisions Made", heading_style))
            for decision in summary.decisions:
                elements.append(Paragraph(f"• {decision}", normal_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Speaker Stats
        if summary.speaker_stats:
            elements.append(Paragraph("💬 Speaking Time", heading_style))
            
            stats_data = [['Speaker', 'Percentage']]
            for speaker, percentage in summary.speaker_stats.items():
                stats_data.append([speaker, f'{percentage:.1f}%'])
            
            stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ]))
            
            elements.append(stats_table)
        
        # Build PDF
        doc.build(elements)
        
        print(f"✅ PDF exported: {filename}")
        return filename
    
    def export_markdown(
        self,
        summary,  # MeetingSummary object
        meeting_name: str
    ) -> Path:
        """
        Export summary as Markdown file
        
        Args:
            summary: MeetingSummary object
            meeting_name: Name of the meeting
        
        Returns:
            Path to generated Markdown file
        """
        print("📝 Generating Markdown...")
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"{meeting_name.replace(' ', '_')}_{timestamp}.md"
        
        # Write markdown content
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(summary.to_markdown())
        
        print(f"✅ Markdown exported: {filename}")
        return filename
    
    def export_docx(
        self,
        summary,  # MeetingSummary object
        meeting_name: str
    ) -> Optional[Path]:
        """
        Export summary as Word document (DOCX)
        
        Args:
            summary: MeetingSummary object
            meeting_name: Name of the meeting
        
        Returns:
            Path to generated DOCX file or None if not available
        """
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx not available. Cannot export to DOCX.")
            return None
        
        print("📄 Generating DOCX...")
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"{meeting_name.replace(' ', '_')}_{timestamp}.docx"
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'🎙️ {meeting_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Date: {summary.timestamp.strftime('%A, %B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Duration: {summary.duration:.1f} minutes")
        
        if summary.participants:
            doc.add_paragraph(f"Participants: {', '.join(summary.participants)}")
        
        doc.add_paragraph()  # Blank line
        
        # Summary
        doc.add_heading('📋 Summary', 1)
        doc.add_paragraph(summary.summary)
        
        # Key Points
        if summary.key_points:
            doc.add_heading('🔑 Key Points', 1)
            for point in summary.key_points:
                doc.add_paragraph(point, style='List Bullet')
        
        # Action Items
        if summary.action_items:
            doc.add_heading('✅ Action Items', 1)
            for item in summary.action_items:
                p = doc.add_paragraph(f'☐ {item}')
                # Highlight action items
                p.runs[0].font.color.rgb = RGBColor(255, 140, 0)
        
        # Decisions
        if summary.decisions:
            doc.add_heading('⚖️ Decisions Made', 1)
            for decision in summary.decisions:
                doc.add_paragraph(decision, style='List Bullet')
        
        # Speaker Stats
        if summary.speaker_stats:
            doc.add_heading('💬 Speaking Time', 1)
            
            # Create table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Speaker'
            header_cells[1].text = 'Percentage'
            
            # Data rows
            for speaker, percentage in summary.speaker_stats.items():
                row_cells = table.add_row().cells
                row_cells[0].text = speaker
                row_cells[1].text = f'{percentage:.1f}%'
        
        # Save document
        doc.save(str(filename))
        # Convert markdown to HTML
        md_content = summary.to_markdown()
        html_content = convert_markdown_to_html(md_content)
    
    def export_html(
        self,
        summary,  # MeetingSummary object
        meeting_name: str
    ) -> Path:
        """
        Export summary as HTML file
        
        Args:
            summary: MeetingSummary object
            meeting_name: Name of the meeting
        
        Returns:
            Path to generated HTML file
        """
        print("🌐 Generating HTML...")
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"{meeting_name.replace(' ', '_')}_{timestamp}.html"
        
        # Convert markdown to HTML
        md_content = summary.to_markdown()
        html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
        
        # Wrap in HTML template
        full_html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{meeting_name} - Meeting Summary</title>
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 900px;
                    margin: 0 auto;
                    padding: 40px 20px;
                    background: #f5f5f5;
                }}
                .container {{
                    background: white;
                    padding: 40px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }}
                h1 {{
                    color: #667eea;
                    border-bottom: 3px solid #667eea;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #764ba2;
                    margin-top: 30px;
                }}
                ul {{
                    padding-left: 25px;
                }}
                li {{
                    margin: 10px 0;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th, td {{
                    padding: 12px;
                    text-align: left;
                    border: 1px solid #ddd;
                }}
                th {{
                    background-color: #667eea;
                    color: white;
                    font-weight: bold;
                }}
                tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                .metadata {{
                    background: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    margin: 20px 0;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                {html_content}
            </div>
        </body>
        </html>
        """
        
        # Write HTML file
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        print(f"✅ HTML exported: {filename}")
        return filename
    
    def export_all_formats(
        self,
        summary,  # MeetingSummary object
        meeting_name: str
    ) -> dict:
        """
        Export summary in all available formats
        
        Returns:
            Dictionary with format names as keys and file paths as values
        """
        print("📦 Exporting to all formats...")
        
        exported_files = {}
        
        try:
            exported_files['pdf'] = self.export_pdf(summary, meeting_name)
        except Exception as e:
            print(f"⚠️ PDF export failed: {e}")
        
        try:
            exported_files['markdown'] = self.export_markdown(summary, meeting_name)
        except Exception as e:
            print(f"⚠️ Markdown export failed: {e}")
        
        try:
            docx_file = self.export_docx(summary, meeting_name)
            if docx_file:
                exported_files['docx'] = docx_file
        except Exception as e:
            print(f"⚠️ DOCX export failed: {e}")
        
        try:
            exported_files['html'] = self.export_html(summary, meeting_name)
        except Exception as e:
            print(f"⚠️ HTML export failed: {e}")
        
        print(f"\n✅ Exported {len(exported_files)} formats successfully!")
        return exported_files


def test_export():
    """Test export functionality"""
    from summarization.summarizer import MeetingSummary
    
    print("🎯 Testing Export Manager")
    print("=" * 50)
    
    # Create sample summary
    sample_summary = MeetingSummary(
        raw_transcript="Sample transcript...",
        summary="This was a productive team meeting where we discussed project progress and next steps.",
        key_points=[
            "Frontend development is 80% complete",
            "API integration scheduled for next week",
            "Team morale is high and everyone is on track"
        ],
        action_items=[
            "John to complete API documentation by Thursday",
            "Sarah to schedule design review meeting",
            "Team to update project timeline in Jira"
        ],
        decisions=[
            "Approved budget increase for additional tooling",
            "Extended sprint by 3 days to accommodate testing",
            "Agreed on new code review process"
        ],
        participants=["John Doe", "Jane Smith", "Bob Johnson"],
        duration=45.5,
        timestamp=datetime.now(),
        speaker_stats={
            "SPEAKER_0": 35.2,
            "SPEAKER_1": 42.8,
            "SPEAKER_2": 22.0
        }
    )
    
    # Create export manager
    export_manager = ExportManager()
    
    # Export to all formats
    files = export_manager.export_all_formats(sample_summary, "Test Team Meeting")
    
    print("\n📁 Exported Files:")
    print("-" * 50)
    for format_name, file_path in files.items():
        print(f"✅ {format_name.upper()}: {file_path}")


if __name__ == "__main__":
    test_export()